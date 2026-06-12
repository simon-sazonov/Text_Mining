#!/usr/bin/env python
"""Build report_12.docx mirroring the PDF: cover page, serif styles, 1.5 spacing,
footer page numbers (cover unnumbered), APA hanging-indent references.
Run from group_12/report/:  ../../../nlp2026/.venv python? -> use the course venv python."""
import subprocess

subprocess.run(["pandoc", "report_12.md", "--number-sections", "-o", "report_12.docx"], check=True)

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document("report_12.docx")

# base + heading styles (theme font references removed so the serif sticks)
normal = doc.styles["Normal"]
normal.font.name = "Cambria"
normal.font.size = Pt(11)
normal.paragraph_format.line_spacing = 1.5
for name, size in [("Heading 1", 15), ("Heading 2", 12.5)]:
    st = doc.styles[name]
    st.font.size = Pt(size)
    st.font.bold = True
    st.font.color.rgb = RGBColor(0, 0, 0)
    st.paragraph_format.space_before = Pt(20)
    st.paragraph_format.space_after = Pt(10)
for name in ("Heading 1", "Heading 2", "Heading 3", "Title"):
    try:
        st = doc.styles[name]
    except KeyError:
        continue
    rPr = st.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    for attr in ("asciiTheme", "hAnsiTheme", "cstheme", "eastAsiaTheme"):
        rFonts.attrib.pop(qn(f"w:{attr}"), None)
    rFonts.set(qn("w:ascii"), "Cambria")
    rFonts.set(qn("w:hAnsi"), "Cambria")

# cover page, mirroring preamble.typ
first = doc.paragraphs[0]._p

def cover_par(text="", size=11, bold=False, lines=None, space_after=4):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    for i, chunk in enumerate(lines or [text]):
        if i:
            p.runs[-1].add_break(WD_BREAK.LINE)
        r = p.add_run(chunk)
        r.font.name = "Cambria"
        r.font.size = Pt(size)
        r.bold = bold
    first.addprevious(p._p)

for _ in range(7):
    cover_par(" ")
cover_par(lines=["Market Sentiment Classification", "of Financial Tweets"], size=22, bold=True, space_after=18)
cover_par("Text Mining — Project Report", 13, space_after=8)
cover_par(lines=["NOVA IMS, Information Management School", "Spring Semester 2025/2026"], size=11.5, space_after=30)
cover_par("Group 12", 12, bold=True, space_after=8)
for member, sid in [("Artem Polikarpov", "20250443"), ("João Cardoso", "20240529"), ("Simon Sazonov", "20221689")]:
    cover_par(f"{member}  {sid}", space_after=2)
cover_par(" ", space_after=24)
cover_par("June 2026")
brk = doc.add_paragraph()
brk.add_run().add_break(WD_BREAK.PAGE)
first.addprevious(brk._p)

# APA hanging indent for everything after the References heading
in_refs = False
for p in doc.paragraphs:
    if p.style.name.startswith("Heading") and p.text.strip() == "References":
        in_refs = True
        continue
    if in_refs and p.text.strip():
        p.paragraph_format.left_indent = Pt(36)
        p.paragraph_format.first_line_indent = Pt(-36)
        p.paragraph_format.line_spacing = 1.5

# footer page numbers; cover unnumbered, body starts at 1
sec = doc.sections[0]
sec.different_first_page_header_footer = True
pg = OxmlElement("w:pgNumType")
pg.set(qn("w:start"), "0")
sec._sectPr.append(pg)
fp = sec.footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fld = OxmlElement("w:fldSimple")
fld.set(qn("w:instr"), "PAGE")
run = OxmlElement("w:r")
t = OxmlElement("w:t")
t.text = "1"
run.append(t)
fld.append(run)
fp._p.append(fld)

doc.save("report_12.docx")
print("report_12.docx built")
